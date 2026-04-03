from __future__ import annotations

import os
from dataclasses import dataclass
from pathlib import Path

from pydantic import BaseModel


class ParsedDocument(BaseModel):
    text: str
    metadata: dict[str, str]


class DocumentParser:
    async def parse_document(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        filename = path.name

        if ext == ".pdf":
            return self._parse_pdf(str(path), filename)
        if ext == ".docx":
            return self._parse_docx(str(path), filename)
        if ext in {".txt", ".md", ".csv", ".log"}:
            return self._parse_txt(str(path), filename)

        # Fallback: treat as text.
        return self._parse_txt(str(path), filename)

    def _parse_pdf(self, file_path: str, filename: str) -> ParsedDocument:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        parts: list[str] = []
        for page in reader.pages:
            try:
                extracted = page.extract_text() or ""
            except Exception:
                extracted = ""
            if extracted.strip():
                parts.append(extracted)
        text = "\n".join(parts).strip()
        return ParsedDocument(
            text=text,
            metadata={"source_filename": filename, "file_type": "pdf"},
        )

    def _parse_docx(self, file_path: str, filename: str) -> ParsedDocument:
        from docx import Document

        doc = Document(file_path)
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        text = "\n".join(parts).strip()
        return ParsedDocument(
            text=text,
            metadata={"source_filename": filename, "file_type": "docx"},
        )

    def _parse_txt(self, file_path: str, filename: str) -> ParsedDocument:
        # Keep parsing deterministic: try utf-8 first, fall back to latin-1.
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()
        return ParsedDocument(
            text=text.strip(),
            metadata={"source_filename": filename, "file_type": "txt"},
        )

